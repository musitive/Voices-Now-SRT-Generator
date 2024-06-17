from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement,qn
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ================================================================================================

# Column indices
LOOP = 0
IN = 1
OUT = 2
ROLE = 3
ENGLISH = 4
TRANSLATION = 5

# Column headers
OLD_LOOP_HEADER = "LOOP"
NEW_LOOP_HEADER = "#"
IN_HEADER = "IN"
OUT_HEADER = "OUT"

# Column widths
LOOP_WIDTH = Inches(0.25)
INTIME_WIDTH = Inches(1)
OUTTIME_WIDTH = Inches(1)
ROLE_WIDTH = Inches(1)
ENGLISH_WIDTH = Inches(2.8)
TRANSLATION_WIDTH = Inches(2.8)

# Default styles
DEFAULT_STYLE = 'Normal'
DEFAULT_FONT = 'Arial'
DEFAULT_FONT_SIZE = Pt(11)
DEFAULT_HEADER_COLOR = 'F6CC9E'

# ================================================================================================

# ------------------------------------------------------------------------------------
# Color a cell
# cell: Cell    - the cell to color
## returns: None
def color_cell(cell) -> None:
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), DEFAULT_HEADER_COLOR))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)
# ------------------------------------------------------------------------------------

# ------------------------------------------------------------------------------------
# Set repeat table row on every new page
# header_row: Row    - the header row to set
## returns: Row
def set_repeat_table_header(header_row):
    tr = header_row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return header_row
# ------------------------------------------------------------------------------------

# ------------------------------------------------------------------------------------
# Set the widths of the cells in the row
# row: Row    - the row to set the widths of
## returns: None
def set_widths(row):
    row.cells[LOOP].width = LOOP_WIDTH
    row.cells[IN].width = INTIME_WIDTH
    row.cells[OUT].width = OUTTIME_WIDTH
    row.cells[ROLE].width = ROLE_WIDTH
    row.cells[ENGLISH].width = ENGLISH_WIDTH
    row.cells[TRANSLATION].width = TRANSLATION_WIDTH
# ------------------------------------------------------------------------------------

# ================================================================================================

class ScriptStyler:
    # --------------------------------------------------------------------------------
    # ScriptStyler
    # document: Document    - the document to style
    def __init__(self, document: Document):
        self.document = document
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Initialize the tables in the document
    ## returns: None
    def initialize_tables(self) -> None:
        self.remove_metadata()

        self.old_table = self.document.tables[0]
        self.new_table = self.document.add_table(1, 6)
        self.new_table.style = 'Table Grid'
        set_repeat_table_header(self.new_table.rows[0])

        self.initialize_table_headers()
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Remove a table from the document
    # index: int    - the index of the table to remove
    ## returns: None
    def remove_table(self, index: int) -> None:
        table = self.document.tables[index]
        table._element.getparent().remove(table._element)
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Remove the metadata from the document
    ## returns: None
    def remove_metadata(self) -> None:
        self.remove_table(2)
        self.remove_table(0)
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Change the orientation of the document to landscape
    ## returns: None
    def change_orientation(self) -> None:
        sections = self.document.sections

        for section in sections:
            # change orientation to landscape
            section.orientation = WD_ORIENT.LANDSCAPE

            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # https://github.com/python-openxml/python-docx/issues/245#event-621236139
    # Prevent table cells from splitting across pages
    ## returns: None
    def prevent_document_break(self) -> None:
        tags = self.document.element.xpath('//w:tr')
        rows = len(tags)
        for row in range(0, rows):
            tag = tags[row]                     # Specify which <w:r> tag you want
            child = OxmlElement('w:cantSplit')  # Create arbitrary tag
            tag.append(child)                   # Append in the new tag
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Save the document as a new file
    # new_filename: str    - the name of the new file
    ## returns: None
    def save_as_new_script(self, new_filename) -> None:
        self.document.save(new_filename)
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Set the styles of the document
    # style: str        - the style to set
    # font_name: str    - the font name to set
    # font_size: str    - the font size to set
    ## returns: None
    def set_styles(self, style: str = DEFAULT_STYLE, font_name: str = DEFAULT_FONT, font_size: str = DEFAULT_FONT_SIZE) -> None:
        s = self.document.styles[style]
        font = s.font
        font.name = font_name
        font.size = font_size
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Initialize the headers of the table
    ## returns: None
    def initialize_table_headers(self) -> None:
        old_row = self.old_table.rows[0]
        new_row = self.new_table.rows[0]

        self.shift_translation(old_row, new_row, IN_HEADER, OUT_HEADER, bold=True, loop=NEW_LOOP_HEADER)
        for cell in new_row.cells:
            color_cell(cell)
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Add a row to the new table
    # translation_index: int    - the index of the translation in the old table
    # in_time: str             - the time the caption appears on screen
    # out_time: str            - the time the caption disappears from the screen
    ## returns: None
    def add_row_to_new_table(self, translation_index: int, in_time: str, out_time: str) -> None:
        old_row = self.old_table.rows[translation_index]
        new_row = self.new_table.add_row()
        self.shift_translation(old_row, new_row, in_time, out_time, loop=str(translation_index))
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Shift the translation from the old row to the new row
    # old_row: Row    - the old row to shift from
    # new_row: Row    - the new row to shift to
    # in_time: str    - the time the caption appears on screen
    # out_time: str   - the time the caption disappears from the screen
    # bold: bool      - whether to bold the text
    # style: str      - the style to set
    # loop: str       - the loop number
    ## returns: None
    def shift_translation(self, old_row, new_row, in_time: str, out_time: str, bold: bool = False, style: str = DEFAULT_STYLE, loop: str = "") -> None:
        if loop == "":
            loop = old_row.cells[self.old_start].text
        self.update_text(new_row, loop, LOOP, bold, style)
        self.update_text(new_row, in_time, IN, bold, style)
        self.update_text(new_row, out_time, OUT, bold, style)
        
        for i in range(1, len(old_row.cells)-self.old_start):
            contents = old_row.cells[i+self.old_start].text
            self.update_text(new_row, contents, i+2, bold, style)

        set_widths(new_row)
    # --------------------------------------------------------------------------------

    # --------------------------------------------------------------------------------
    # Update the text in a cell
    # new_row: Row    - the row to update
    # text: str       - the text to add
    # cell: int       - the cell to update
    # bold: bool      - whether to bold the text
    # style: str      - the style to set
    ## returns: None
    def update_text(self, new_row, text: str, cell: int, bold: bool = False, style: str = DEFAULT_STYLE):
        run = new_row.cells[cell].paragraphs[0].add_run(text)
        new_row.cells[cell].paragraphs[0].style = self.document.styles[style]
        run.bold = bold
    # --------------------------------------------------------------------------------

# ================================================================================================
