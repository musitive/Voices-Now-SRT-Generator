from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement,qn
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def color_cell(cell) -> None:
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="F6CC9E"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

def set_repeat_table_header(header_row):
    """ set repeat table row on every new page
    """
    tr = header_row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return header_row

def set_widths(row):
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(1)
    row.cells[2].width = Inches(1.5)
    row.cells[3].width = Inches(3)
    row.cells[4].width = Inches(3)

class LdsScript:
    def __init__(self, wordFilename: str):
        self.document = Document(wordFilename)

        # Python-docx set-up
        # Most LDS Translations should be in the second table, fourth column
        self.tables = self.document.tables
        self.columns = self.tables[1].columns
        if len(self.columns) == 4:
            self.translation = self.columns[3]
        elif len(self.columns) == 5:
            self.translation = self.columns[4]
        else:
            raise Exception("Unable to resolve this LDS Script format.")
        self.newSection = True

    def get_translation(self, translationRow: int) -> str:
        return self.translation.cells[translationRow].text
    
    def change_orientation(self) -> None:
        sections = self.document.sections

        for section in sections:
            # change orientation to landscape
            section.orientation = WD_ORIENT.LANDSCAPE

            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

    def remove_table(self, index: int) -> None:
        table = self.document.tables[index]
        table._element.getparent().remove(table._element)

    def remove_metadata(self) -> None:
        self.remove_table(2)
        self.remove_table(0)

    def initialize_tables(self) -> None:
        self.remove_metadata()

        self.old_table = self.document.tables[0]
        self.new_table = self.document.add_table(1, 5)
        self.new_table.style = 'Table Grid'
        set_repeat_table_header(self.new_table.rows[0])

        self.initialize_table_headers()

    def initialize_table_headers(self) -> None:
        old_row = self.old_table.rows[0]
        new_row = self.new_table.rows[0]

        self.shift_translation(old_row, new_row, "IN", True)
        for cell in new_row.cells:
            color_cell(cell)
    
    def prevent_document_break(self) -> None:
        """https://github.com/python-openxml/python-docx/issues/245#event-621236139
        Globally prevent table cells from splitting across pages.
        """
        tags = self.document.element.xpath('//w:tr')
        rows = len(tags)
        for row in range(0, rows):
            tag = tags[row]  # Specify which <w:r> tag you want
            child = OxmlElement('w:cantSplit')  # Create arbitrary tag
            tag.append(child)  # Append in the new tag

    def add_row_to_new_table(self, translation_index: int, in_time: str) -> None:
        old_row = self.old_table.rows[translation_index]
        new_row = self.new_table.add_row()
        self.shift_translation(old_row, new_row, in_time)

    def shift_translation(self, old_row, new_row, in_time: str, bold: bool = False, style: str = 'Normal') -> None:
        contents = old_row.cells[0].text
        self.update_text(new_row, contents, 0, bold, style)
        self.update_text(new_row, in_time, 1, bold, style)
        
        for n in range(1, len(old_row.cells)):
            contents = old_row.cells[n].text
            self.update_text(new_row, contents, n+1, bold, style)

        set_widths(new_row)
        return

    def update_text(self, new_row, text: str, cell: int, bold: bool = False, style: str = 'Normal', ):
        run = new_row.cells[cell].paragraphs[0].add_run(text)
        new_row.cells[cell].paragraphs[0].style = self.document.styles[style]
        run.bold = bold

    def save_as_new_script(self, new_filename) -> None:
        self.document.save(new_filename)

    def set_styles(self) -> None:
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)